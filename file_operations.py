import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import logging

logger = logging.getLogger(__name__)


class FileOperations:
    @staticmethod
    def new_document():
        return ""
    
    @staticmethod
    def open_txt(filepath: str) -> str:
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error opening TXT file: {e}")
            raise
    
    @staticmethod
    def open_docx(filepath: str) -> str:
        try:
            doc = Document(filepath)
            text_content = []
            for para in doc.paragraphs:
                text_content.append(para.text)
            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error opening DOCX file: {e}")
            raise
    
    @staticmethod
    def save_txt(filepath: str, content: str):
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            logger.info(f"File saved: {filepath}")
        except Exception as e:
            logger.error(f"Error saving TXT file: {e}")
            raise
    
    @staticmethod
    def save_docx(filepath: str, content: str, formatting: dict = None):
        try:
            doc = Document()
            
            if formatting is None:
                formatting = {}
            
            paragraphs = content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text)
                    
                    if 'font_name' in formatting:
                        para.runs[0].font.name = formatting['font_name']
                    if 'font_size' in formatting:
                        para.runs[0].font.size = Pt(formatting['font_size'])
                    if 'bold' in formatting:
                        para.runs[0].font.bold = formatting['bold']
                    if 'italic' in formatting:
                        para.runs[0].font.italic = formatting['italic']
                    if 'underline' in formatting:
                        para.runs[0].font.underline = formatting['underline']
                else:
                    doc.add_paragraph('')
            
            doc.save(filepath)
            logger.info(f"DOCX file saved: {filepath}")
        except Exception as e:
            logger.error(f"Error saving DOCX file: {e}")
            raise
    
    @staticmethod
    def export_pdf(filepath: str, content: str, formatting: dict = None):
        try:
            c = canvas.Canvas(filepath, pagesize=letter)
            width, height = letter
            
            if formatting is None:
                formatting = {}
            
            font_size = formatting.get('font_size', 12)
            
            y_position = height - 1 * inch
            line_height = font_size + 4
            
            lines = content.split('\n')
            for line in lines:
                if y_position < 1 * inch:
                    c.showPage()
                    y_position = height - 1 * inch
                
                c.setFont("Helvetica", font_size)
                
                max_width = width - 2 * inch
                words = line.split()
                current_line = ""
                
                for word in words:
                    test_line = current_line + " " + word if current_line else word
                    if c.stringWidth(test_line, "Helvetica", font_size) <= max_width:
                        current_line = test_line
                    else:
                        if current_line:
                            c.drawString(1 * inch, y_position, current_line)
                            y_position -= line_height
                            if y_position < 1 * inch:
                                c.showPage()
                                y_position = height - 1 * inch
                        current_line = word
                
                if current_line:
                    c.drawString(1 * inch, y_position, current_line)
                    y_position -= line_height
            
            c.save()
            logger.info(f"PDF file exported: {filepath}")
        except Exception as e:
            logger.error(f"Error exporting PDF: {e}")
            raise
    
    @staticmethod
    def get_file_extension(filepath: str) -> str:
        return os.path.splitext(filepath)[1].lower()
    
    @staticmethod
    def is_valid_path(filepath: str) -> bool:
        directory = os.path.dirname(filepath)
        return os.path.exists(directory) if directory else True
